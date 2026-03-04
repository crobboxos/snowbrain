import json
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

from docx import Document


ROOT = Path("generated")
REPORTS = ROOT / "REPORTS"
OUT_DOCX = ROOT / "SERVICENOW_CUSTOM_ASPECTS_AUDIT.docx"
INTEGRATION_MARKERS = ("xeretec", "freshdesk", "dev.azure.com", "salesforce.com")


def read_json(path):
    return json.loads(path.read_text(encoding="utf-8"))


def text(value):
    return "" if value is None else str(value).strip()


def is_true(value):
    return text(value).lower() in {"1", "true", "yes", "y"}


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    head_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        head_cells[index].text = text(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = text(value)
    document.add_paragraph("")


def scope_display(scope_id, scope_map):
    scope_id = text(scope_id)
    if not scope_id:
        return "Global"
    if scope_id.lower() == "global":
        return "Global"
    scope = scope_map.get(scope_id)
    if not scope:
        return scope_id
    code = text(scope.get("scope"))
    name = text(scope.get("name"))
    if code and name:
        return f"{code} ({name})"
    return code or name or scope_id


def parse_acl_table(name):
    acl_name = text(name)
    if not acl_name:
        return ""
    if "." in acl_name:
        return acl_name.split(".", 1)[0]
    return acl_name


def is_custom_integration(name, endpoint):
    normalized = f"{text(name).lower()} {text(endpoint).lower()}"
    return any(marker in normalized for marker in INTEGRATION_MARKERS)


def main():
    focus_path = REPORTS / "custom_focus.json"
    if not focus_path.exists():
        raise SystemExit(
            "Missing generated/REPORTS/custom_focus.json. Run the custom extraction step first."
        )

    focus = read_json(focus_path)
    data = focus.get("data", {})

    scope_inventory_path = REPORTS / "scope_inventory.json"
    if scope_inventory_path.exists():
        scope_inventory = read_json(scope_inventory_path)
    else:
        scope_inventory = data.get("sys_scope", [])

    scope_map = {text(scope.get("sys_id")): scope for scope in scope_inventory if text(scope.get("sys_id"))}

    tables = data.get("sys_db_object_custom", [])
    dictionary_rows = data.get("sys_dictionary_custom", [])
    business_rules = data.get("sys_script_custom", [])
    script_includes = data.get("sys_script_include_custom", [])
    client_scripts = data.get("sys_script_client_custom", [])
    ui_actions = data.get("sys_ui_action_custom", [])
    scheduled_jobs = data.get("sysauto_script_custom", [])
    flows = data.get("sys_hub_flow_custom", [])
    rest_messages = data.get("sys_rest_message_custom", [])
    rest_methods = data.get("sys_rest_message_fn_custom", [])
    soap_messages = data.get("sys_soap_message_custom", [])
    acl_rows = data.get("sys_security_acl_custom", [])

    dependency_edges_path = REPORTS / "dependency_edges.json"
    dependency_edges = read_json(dependency_edges_path) if dependency_edges_path.exists() else []

    field_counts = Counter()
    reference_counts = Counter()
    mandatory_counts = Counter()
    for row in dictionary_rows:
        table_name = text(row.get("name"))
        if not table_name:
            continue
        field_counts[table_name] += 1
        if text(row.get("reference")):
            reference_counts[table_name] += 1
        if is_true(row.get("mandatory")):
            mandatory_counts[table_name] += 1

    acl_counts = Counter()
    acl_operation_counts = Counter()
    for acl in acl_rows:
        table_name = parse_acl_table(acl.get("name"))
        if table_name:
            acl_counts[table_name] += 1
        acl_operation_counts[text(acl.get("operation"))] += 1

    br_table_counts = Counter()
    for rule in business_rules:
        table_name = text(rule.get("collection"))
        if table_name:
            br_table_counts[table_name] += 1

    table_rows = []
    for table in tables:
        table_name = text(table.get("name"))
        risk = (
            field_counts[table_name]
            + (reference_counts[table_name] * 3)
            + (mandatory_counts[table_name] * 2)
            + (acl_counts[table_name] * 4)
            + (br_table_counts[table_name] * 5)
        )
        table_rows.append(
            {
                "table": table_name,
                "label": text(table.get("label")),
                "scope_id": text(table.get("sys_scope")) or "global",
                "created_by": text(table.get("sys_created_by")),
                "updated_on": text(table.get("sys_updated_on")),
                "fields": field_counts[table_name],
                "references": reference_counts[table_name],
                "mandatory": mandatory_counts[table_name],
                "acls": acl_counts[table_name],
                "business_rules": br_table_counts[table_name],
                "risk": risk,
            }
        )
    table_rows.sort(key=lambda row: row["risk"], reverse=True)

    process_artifacts = []
    for row in business_rules:
        process_artifacts.append(
            {
                "type": "Business Rule",
                "name": text(row.get("name")),
                "table": text(row.get("collection")),
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "created_by": text(row.get("sys_created_by")),
                "updated_on": text(row.get("sys_updated_on")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    for row in script_includes:
        process_artifacts.append(
            {
                "type": "Script Include",
                "name": text(row.get("name")),
                "table": "",
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "created_by": text(row.get("sys_created_by")),
                "updated_on": text(row.get("sys_updated_on")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    for row in client_scripts:
        process_artifacts.append(
            {
                "type": "Client Script",
                "name": text(row.get("name")),
                "table": text(row.get("table")),
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "created_by": text(row.get("sys_created_by")),
                "updated_on": text(row.get("sys_updated_on")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    for row in ui_actions:
        process_artifacts.append(
            {
                "type": "UI Action",
                "name": text(row.get("name")),
                "table": text(row.get("table")),
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "created_by": text(row.get("sys_created_by")),
                "updated_on": text(row.get("sys_updated_on")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    for row in scheduled_jobs:
        process_artifacts.append(
            {
                "type": "Scheduled Script",
                "name": text(row.get("name")),
                "table": "",
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "created_by": text(row.get("sys_created_by")),
                "updated_on": text(row.get("sys_updated_on")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    for row in flows:
        process_artifacts.append(
            {
                "type": "Flow",
                "name": text(row.get("name")),
                "table": "",
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "created_by": text(row.get("sys_created_by")),
                "updated_on": text(row.get("sys_updated_on")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    process_artifacts.sort(key=lambda row: (row["type"], row["name"], row["sys_id"]))

    rest_message_names = {text(row.get("sys_id")): text(row.get("name")) for row in rest_messages}

    integration_rows = []
    host_counts = Counter()
    for row in rest_messages:
        endpoint = text(row.get("rest_endpoint"))
        if not is_custom_integration(row.get("name"), endpoint):
            continue
        host = text(urlparse(endpoint).hostname)
        if host:
            host_counts[host] += 1
        integration_rows.append(
            {
                "type": "REST Message",
                "name": text(row.get("name")),
                "method": "",
                "endpoint": endpoint,
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    for row in rest_methods:
        endpoint = text(row.get("rest_endpoint"))
        parent = rest_message_names.get(text(row.get("rest_message")))
        if not is_custom_integration(parent or row.get("name"), endpoint):
            continue
        host = text(urlparse(endpoint).hostname)
        if host:
            host_counts[host] += 1
        integration_rows.append(
            {
                "type": "REST Method",
                "name": parent or text(row.get("name")),
                "method": text(row.get("http_method")).upper(),
                "endpoint": endpoint,
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    for row in soap_messages:
        endpoint = text(row.get("endpoint"))
        if not is_custom_integration(row.get("name"), endpoint):
            continue
        host = text(urlparse(endpoint).hostname)
        if host:
            host_counts[host] += 1
        integration_rows.append(
            {
                "type": "SOAP Message",
                "name": text(row.get("name")),
                "method": "SOAP",
                "endpoint": endpoint,
                "scope_id": text(row.get("sys_scope")),
                "active": is_true(row.get("active")),
                "sys_id": text(row.get("sys_id")),
            }
        )
    integration_rows.sort(key=lambda row: (row["type"], row["name"], row["sys_id"]))

    custom_artifact_ids = {artifact["sys_id"] for artifact in process_artifacts if artifact["sys_id"]}
    custom_dependencies = [
        edge for edge in dependency_edges if text(edge.get("artifact_sys_id")) in custom_artifact_ids
    ]
    custom_dependencies.sort(key=lambda edge: (text(edge.get("source")), text(edge.get("target"))))

    scope_usage = defaultdict(lambda: {"tables": 0, "process": 0, "integrations": 0, "acls": 0})
    for row in tables:
        scope_id = text(row.get("sys_scope"))
        if scope_id and scope_id.lower() != "global":
            scope_usage[scope_id]["tables"] += 1
    for row in process_artifacts:
        scope_id = text(row.get("scope_id"))
        if scope_id and scope_id.lower() != "global":
            scope_usage[scope_id]["process"] += 1
    for row in integration_rows:
        scope_id = text(row.get("scope_id"))
        if scope_id and scope_id.lower() != "global":
            scope_usage[scope_id]["integrations"] += 1
    for row in acl_rows:
        scope_id = text(row.get("sys_scope"))
        if scope_id and scope_id.lower() != "global":
            scope_usage[scope_id]["acls"] += 1

    scope_rows = []
    for scope_id, usage in sorted(
        scope_usage.items(),
        key=lambda item: (item[1]["tables"] + item[1]["process"] + item[1]["integrations"] + item[1]["acls"]),
        reverse=True,
    ):
        scope = scope_map.get(scope_id, {})
        scope_rows.append(
            [
                scope_id,
                text(scope.get("scope")),
                text(scope.get("name")),
                text(scope.get("vendor")) or "Unknown",
                text(scope.get("private")),
                usage["tables"],
                usage["process"],
                usage["integrations"],
                usage["acls"],
            ]
        )

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    document = Document()
    document.add_heading("ServiceNow Custom Aspects Deep-Dive", 0)
    document.add_paragraph(f"Generated: {generated_at}")
    document.add_paragraph(f"Instance: {text(focus.get('base_url'))}")
    document.add_paragraph(
        "This report documents custom and custom-adjacent ServiceNow assets discovered from live API queries."
    )

    document.add_heading("1. Executive Summary", level=1)
    summary_points = [
        f"Custom table records discovered: {len(tables)}",
        f"Dictionary rows tied to custom tables/fields: {len(dictionary_rows)}",
        f"Custom process artifacts discovered: {len(process_artifacts)}",
        f"Custom integration artifacts discovered: {len(integration_rows)}",
        f"Custom ACL records discovered: {len(acl_rows)}",
        f"Dependency edges linked to custom artifacts: {len(custom_dependencies)}",
    ]
    for point in summary_points:
        document.add_paragraph(point, style="List Bullet")

    document.add_heading("2. Identification Rules", level=1)
    for point in [
        "Table and artifact names starting with u_ or x_.",
        "Artifacts containing business keywords xeretec or freshdesk.",
        "Artifacts linked to custom tables through table/collection fields.",
        "Scope IDs resolved to scope code and app name where possible.",
        "Dependency edges filtered to artifacts discovered by the custom queries.",
    ]:
        document.add_paragraph(point, style="List Bullet")

    document.add_heading("3. Scope Register", level=1)
    if scope_rows:
        add_table(
            document,
            ["Scope Sys ID", "Scope Code", "App Name", "Vendor", "Private", "Tables", "Process", "Integrations", "ACLs"],
            scope_rows,
        )
    else:
        document.add_paragraph("No non-global scopes were detected in the custom dataset.")

    document.add_heading("4. Custom Data Model Register", level=1)
    document.add_paragraph("Tables are sorted by derived risk score (fields, references, mandatory flags, ACLs, and BR count).")
    add_table(
        document,
        ["Table", "Label", "Scope", "Created By", "Fields", "Ref Fields", "Mandatory", "ACLs", "BRs", "Risk"],
        [
            [
                row["table"],
                row["label"],
                scope_display(row["scope_id"], scope_map),
                row["created_by"],
                row["fields"],
                row["references"],
                row["mandatory"],
                row["acls"],
                row["business_rules"],
                row["risk"],
            ]
            for row in table_rows
        ],
    )

    document.add_heading("5. Custom Process Automation", level=1)
    process_type_counter = Counter(row["type"] for row in process_artifacts)
    for process_type, count in sorted(process_type_counter.items(), key=lambda item: (-item[1], item[0])):
        document.add_paragraph(f"{process_type}: {count}", style="List Bullet")
    add_table(
        document,
        ["Type", "Name", "Table", "Scope", "Active", "Created By", "Updated On", "Sys ID"],
        [
            [
                row["type"],
                row["name"],
                row["table"],
                scope_display(row["scope_id"], scope_map),
                "true" if row["active"] else "false",
                row["created_by"],
                row["updated_on"],
                row["sys_id"],
            ]
            for row in process_artifacts
        ],
    )

    document.add_heading("6. Custom Integrations", level=1)
    for host, count in host_counts.most_common():
        document.add_paragraph(f"{host}: {count}", style="List Bullet")
    add_table(
        document,
        ["Type", "Name", "Method", "Endpoint", "Scope", "Active", "Sys ID"],
        [
            [
                row["type"],
                row["name"],
                row["method"],
                row["endpoint"],
                scope_display(row["scope_id"], scope_map),
                "true" if row["active"] else "false",
                row["sys_id"],
            ]
            for row in integration_rows
        ],
    )

    document.add_heading("7. Custom ACL Matrix", level=1)
    for operation, count in acl_operation_counts.most_common():
        document.add_paragraph(f"{operation}: {count}", style="List Bullet")
    add_table(
        document,
        ["Name", "Operation", "Type", "Scope", "Requires Role", "Active", "Created By", "Updated On", "Sys ID"],
        [
            [
                text(row.get("name")),
                text(row.get("operation")),
                text(row.get("type")),
                scope_display(text(row.get("sys_scope")), scope_map),
                text(row.get("requires_role")),
                "true" if is_true(row.get("active")) else "false",
                text(row.get("sys_created_by")),
                text(row.get("sys_updated_on")),
                text(row.get("sys_id")),
            ]
            for row in sorted(
                acl_rows,
                key=lambda acl: (text(acl.get("name")), text(acl.get("operation")), text(acl.get("sys_id"))),
            )
        ],
    )

    document.add_heading("8. Custom Dependency Chains", level=1)
    if custom_dependencies:
        add_table(
            document,
            ["Source", "Target", "Evidence", "Artifact Sys ID"],
            [
                [
                    text(edge.get("source")),
                    text(edge.get("target")),
                    text(edge.get("evidence"))[:160],
                    text(edge.get("artifact_sys_id")),
                ]
                for edge in custom_dependencies
            ],
        )
    else:
        document.add_paragraph("No dependency edges were linked to the custom artifact set.")

    document.add_heading("9. Source Artifacts", level=1)
    for item in [
        "generated/REPORTS/custom_focus.json",
        "generated/REPORTS/scope_inventory.json",
        "generated/REPORTS/dependency_edges.json",
        "generated/README.md",
    ]:
        document.add_paragraph(item, style="List Bullet")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
